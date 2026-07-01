#!/usr/bin/env python3
"""Convert the AI integration blog article to a Word document."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# Configure default font
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# Configure heading styles
for level, size in [(1, 22), (2, 16), (3, 13)]:
    hs = doc.styles[f'Heading {level}']
    hs.font.name = '微软雅黑'
    hs.font.size = Pt(size)
    hs.font.bold = True
    hs.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    hs.paragraph_format.space_before = Pt(12)
    hs.paragraph_format.space_after = Pt(6)

# Title
title = doc.add_heading('WF AI中转站 大模型对接完整教程', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x0d, 0x47, 0xa1)

# Subtitle
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('一文搞定所有主流软件的 AI 接口配置')
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
r.italic = True

doc.add_paragraph()  # spacer

def add_table(doc, rows):
    """Add a formatted table."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = 'Light Grid Accent 1'
    for i, (key, val) in enumerate(rows):
        row = table.rows[i]
        row.cells[0].text = key
        row.cells[1].text = val
        # Bold the first column
        for p in row.cells[0].paragraphs:
            for run in p.runs:
                run.bold = True
    return table

def add_code_block(doc, code_text):
    """Add a code block paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x2d, 0x2d, 0x2d)
    # Add a light grey background (shading)
    shading = p._element.get_or_add_pPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): 'F5F5F5'
    })
    shading.append(shd)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.5 * (level + 1))
    return p

def add_note(doc, label, text):
    """Add a highlighted note."""
    p = doc.add_paragraph()
    r = p.add_run(f'{label} ')
    r.font.bold = True
    r.font.color.rgb = RGBColor(0xcc, 0x66, 0x00)
    r = p.add_run(text)
    r.font.size = Pt(11)
    return p

def add_separator(doc):
    p = doc.add_paragraph('─' * 50)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.color.rgb = RGBColor(0xcc, 0xcc, 0xcc)
        run.font.size = Pt(8)

# ============================================================
# Section: API Info
# ============================================================
doc.add_heading('API 信息速览', level=1)

add_table(doc, [
    ('中转站地址', 'https://api.wf1.one'),
    ('API BaseURL', 'https://api.wf1.one/v1'),
    ('额度购买', 'shop.aklibk.com（自动发货，10元=20刀，1倍率）'),
    ('模型与价格', 'api.wf1.one/pricing'),
    ('支持协议', 'OpenAI 兼容接口（Chat Completions API）'),
])

doc.add_paragraph()

# ============================================================
# Section 1: Register & Login
# ============================================================
doc.add_heading('一、注册与登录', level=1)

doc.add_heading('Step 1 — 打开网址', level=2)
p = doc.add_paragraph('浏览器访问 ')
r = p.add_run('api.wf1.one')
r.font.color.rgb = RGBColor(0x0d, 0x47, 0xa1)

doc.add_heading('Step 2 — 注册账户', level=2)
add_bullet(doc, '点击「Sign up」进入注册页')
add_bullet(doc, '填写 用户名 + 邮箱 + 密码 → 提交注册')

add_note(doc, '⚠️ 注意：', '账户名和密码务必牢记，忘记后无法找回！建议将注册信息保存到密码管理器。')

doc.add_heading('Step 3 — 登录', level=2)
doc.add_paragraph('使用刚注册的用户名和密码登录系统。')

# ============================================================
# Section 2: Get API Token
# ============================================================
doc.add_heading('二、获取 API 令牌', level=1)

doc.add_heading('Step 1 — 兑换额度', level=2)
add_bullet(doc, '点击左侧菜单「钱包」→「充值」')
add_bullet(doc, '前往 shop.aklibk.com 购买额度，购买后自动发货兑换码')
add_bullet(doc, '输入兑换码，点击「兑换」，额度即时到账')

doc.add_heading('Step 2 — 创建令牌', level=2)
doc.add_paragraph('点击「令牌管理」→「添加令牌」，填写以下参数：')

add_table(doc, [
    ('参数', '说明'),
    ('名称', '随意填写，例如「我的主力令牌」'),
    ('令牌分组', '默认分组'),
    ('过期时间', '建议选择「永不过期」，也可设置具体时长'),
    ('新建数量', '需要几个令牌就填几（通常1个即可）'),
    ('额度设置', '每个令牌的额度上限。勾选「无限额度」则此设置自动失效'),
])

doc.add_paragraph()
add_note(doc, '💡 提示：', '创建后请立即复制令牌字符串（sk-xxx），关闭页面后将无法再次查看完整令牌！')

# ============================================================
# Section 3: General Config
# ============================================================
doc.add_heading('三、通用配置说明', level=1)
doc.add_paragraph('所有接入本中转站的软件只需要配置以下三个通用参数：')

p = doc.add_paragraph()
r = p.add_run('📎 API BaseURL（接口地址）')
r.bold = True
doc.add_paragraph('https://api.wf1.one/v1').runs[0].font.name = 'Consolas'

p = doc.add_paragraph()
r = p.add_run('🔑 API Key（密钥）')
r.bold = True
doc.add_paragraph('你在令牌管理中创建的 sk-xxx 字符串')

p = doc.add_paragraph()
r = p.add_run('🧠 Model（模型名称）')
r.bold = True
doc.add_paragraph('常用模型：gpt-4o · claude-3.5-sonnet · gemini-pro · deepseek-chat · Qwen2.5-72B 等')

# ============================================================
# Section 4: Software Tutorials
# ============================================================
doc.add_heading('四、各大主流软件对接教程', level=1)

# --- 1. ChatBox ---
doc.add_heading('1. ChatBox — 全平台 AI 聊天客户端', level=2)
p = doc.add_paragraph()
r = p.add_run('支持 Windows / macOS / Linux / iOS / Android')
r.italic = True
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

add_bullet(doc, '打开 ChatBox → ⚙️ 设置 → 模型提供商')
add_bullet(doc, '选择「OpenAI API」或「自定义」')
add_bullet(doc, '填写以下参数：')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.5)
p.add_run('API 域名：').bold = True
p.add_run('https://api.wf1.one/v1')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.5)
p.add_run('API Key：').bold = True
p.add_run('粘贴你的 sk-xxx 令牌')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.5)
p.add_run('模型：').bold = True
p.add_run('手动输入模型名称（如 gpt-4o）')
add_bullet(doc, '保存并开始聊天 🎉')

# --- 2. LobeChat ---
doc.add_heading('2. LobeChat — 现代化 AI 聊天框架', level=2)
doc.add_paragraph('支持插件、知识库、多模态，可自部署')
add_bullet(doc, '打开 LobeChat → 设置 → 语言模型')
add_bullet(doc, '提供商选择「OpenAI」')
add_bullet(doc, '自定义 API 地址：https://api.wf1.one/v1')
add_bullet(doc, 'API Key：粘贴 sk-xxx 令牌')
add_bullet(doc, '默认模型：选择或输入你想要的模型')
add_bullet(doc, '高级设置中可开启「自定义模型列表」→ 自行添加需要的模型')

# --- 3. Cursor ---
doc.add_heading('3. Cursor — AI 代码编辑器', level=2)
add_bullet(doc, '打开 Cursor → Settings → Models')
add_bullet(doc, '关闭默认的 OpenAI 模型')
add_bullet(doc, '开启「Override OpenAI Base URL」')
add_bullet(doc, '填写：')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.5)
p.add_run('Base URL：https://api.wf1.one/v1')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.5)
p.add_run('Default Model：').bold = True
p.add_run('输入模型名（如 claude-3.5-sonnet-20241022 或 deepseek-coder）')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.5)
p.add_run('API Key：').bold = True
p.add_run('粘贴 sk-xxx 令牌')
add_bullet(doc, '在 Model Names 中手动添加你需要的模型名称')
add_bullet(doc, '重启 Cursor 生效')

# --- 4. Windsurf ---
doc.add_heading('4. Windsurf — AI 驱动 IDE', level=2)
add_bullet(doc, '打开 Windsurf → 设置（Ctrl+Shift+P → Open Settings JSON）')
add_bullet(doc, '在 settings.json 中添加：')

add_code_block(doc, '''{
  "windsurf.apiBaseUrl": "https://api.wf1.one/v1",
  "windsurf.apiKey": "你的 sk-xxx 令牌",
  "windsurf.model": "claude-3.5-sonnet"
}''')
add_bullet(doc, '保存后重启 Windsurf')

# --- 5. VS Code + Continue ---
doc.add_heading('5. VS Code + Continue 插件', level=2)
doc.add_paragraph('VS Code 上最流行的 AI 编程助手插件')
add_bullet(doc, '安装 Continue 插件（Marketplace 搜索 Continue）')
add_bullet(doc, '打开 Continue 侧边栏 → ⚙️ 设置')
add_bullet(doc, '编辑配置文件 ~/.continue/config.json')
add_bullet(doc, '添加以下配置：')

add_code_block(doc, '''{
  "models": [{
    "title": "WF AI",
    "provider": "openai",
    "model": "gpt-4o",
    "apiBase": "https://api.wf1.one/v1",
    "apiKey": "***"
  }],
  "tabAutocompleteModel": {
    "title": "WF AutoComplete",
    "provider": "openai",
    "model": "deepseek-coder",
    "apiBase": "https://api.wf1.one/v1",
    "apiKey": "***"
  }
}''')
add_bullet(doc, '保存后重新加载 Continue')

# --- 6. Python ---
doc.add_heading('6. Python — openai 库调用', level=2)
doc.add_paragraph('使用 Python 的官方 SDK 接入')
p = doc.add_paragraph()
p.add_run('安装依赖：').bold = True
add_code_block(doc, 'pip install openai')
p = doc.add_paragraph()
p.add_run('示例代码：').bold = True
add_code_block(doc, '''from openai import OpenAI

client = OpenAI(
    base_url="https://api.wf1.one/v1",
    api_key="你的 sk-xxx 令牌"
)

response = client.chat.completions.create(
    model="gpt-4o",
    messages=[
        {"role": "system", "content": "你是一个助手"},
        {"role": "user", "content": "你好！"}
    ],
    stream=True
)

for chunk in response:
    print(chunk.choices[0].delta.content or "", end="")''')

# --- 7. curl ---
doc.add_heading('7. Shell / curl — 快速验证令牌', level=2)
add_code_block(doc, '''# 测试通话
curl https://api.wf1.one/v1/chat/completions \\
  -H "Content-Type: application/json" \\
  -H "Authorization: Bearer ***" \\
  -d '{
    "model": "gpt-4o",
    "messages": [{"role": "user", "content": "你好，请用一句话介绍自己"}],
    "stream": true
  }'
# 检查服务
curl https://api.wf1.one/v1/models \\
  -H "Authorization: Bearer ***"''')

# --- 8. ChatGPT App ---
doc.add_heading('8. ChatGPT App — 自定义 API', level=2)
doc.add_paragraph('ChatGPT iOS/Android 官方应用支持自定义 API')
add_bullet(doc, '打开 ChatGPT App → ⚙️ Settings')
add_bullet(doc, '选择「OpenAI API」')
add_bullet(doc, 'API URL：https://api.wf1.one/v1')
add_bullet(doc, 'API Key：粘贴 sk-xxx 令牌')
add_bullet(doc, '模型：手动输入（如 gpt-4o）')
doc.add_paragraph('第三方客户端如 OpenCat、PocketPal AI 配置方式类似。')

# --- 9. One API / New-API ---
doc.add_heading('9. One API / New-API — 下级转发', level=2)
doc.add_paragraph('如果你自己部署了 One API 或 New-API，可将本中转站添加为渠道：')
add_bullet(doc, '登录你的 One API 管理后台')
add_bullet(doc, '渠道 → 添加渠道')
add_bullet(doc, '填写：')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.5)
p.add_run('类型：').bold = True
p.add_run('OpenAI')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.5)
p.add_run('名称：').bold = True
p.add_run('WF AI中转站')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.5)
p.add_run('密钥：').bold = True
p.add_run('你的 sk-xxx 令牌')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.5)
p.add_run('Base URL：').bold = True
p.add_run('https://api.wf1.one/v1')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.5)
p.add_run('模型：').bold = True
p.add_run('填写需要的模型名（多个用逗号分隔）')
add_bullet(doc, '保存后即可在下游分发')

# --- 10. Open WebUI ---
doc.add_heading('10. Open WebUI — 自托管 AI 前端', level=2)
add_bullet(doc, '打开 Open WebUI → 管理员面板 → 设置 → 外部连接')
add_bullet(doc, 'OpenAI API 配置：')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.5)
p.add_run('API Base URL：https://api.wf1.one/v1')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.5)
p.add_run('API Key：').bold = True
p.add_run('粘贴 sk-xxx 令牌')
add_bullet(doc, '在模型管理中添加你需要的模型')
add_bullet(doc, '保存并刷新页面')

# --- 11. NextChat ---
doc.add_heading('11. NextChat（ChatGPT-Next-Web）', level=2)
doc.add_paragraph('流行的自部署 AI 聊天前端，Vercel / Docker 均可')
p = doc.add_paragraph()
p.add_run('环境变量配置（Vercel / Docker）：').bold = True
add_code_block(doc, '''OPENAI_API_KEY=*** 你的 sk-xxx 令牌
BASE_URL=https://api.wf1.one
DEFAULT_MODEL=gpt-4o''')
p = doc.add_paragraph()
p.add_run('Docker 部署示例：').bold = True
add_code_block(doc, '''docker run -d -p 3000:3000 \\
  -e OPENAI_API_KEY=*** \\
  -e BASE_URL=https://api.wf1.one \\
  -e DEFAULT_MODEL=gpt-4o \\
  yidadaa/chatgpt-next-web''')

# ============================================================
# Section 5: Quick Verify
# ============================================================
doc.add_heading('五、快速验证', level=1)

p = doc.add_paragraph()
p.add_run('检查服务是否正常').bold = True
add_code_block(doc, '''curl https://api.wf1.one/v1/models \\
  -H "Authorization: Bearer ***"''')
doc.add_paragraph('返回 JSON 模型列表即表示服务正常 ✅')

p = doc.add_paragraph()
p.add_run('列出所有模型（Python）').bold = True
add_code_block(doc, '''from openai import OpenAI
client = OpenAI(
    base_url="https://api.wf1.one/v1",
    api_key="你的 sk-xxx"
)
models = client.models.list()
for m in models:
    print(m.id)''')

# ============================================================
# Section 6: FAQ
# ============================================================
doc.add_heading('六、常见问题', level=1)

faqs = [
    ('令牌创建后看不到了怎么办？', '关闭页面后无法再次查看完整的 sk-xxx。请在令牌管理中删除原令牌，重新创建一个新的。'),
    ('提示「Insufficient quota」或额度不足？', '登录 api.wf1.one → 钱包查看余额。前往 shop.aklibk.com 购买补充。'),
    ('调用时报错「Invalid token」？', '检查 API Key 是否完整复制（不含多余空格），确认令牌未过期或被删除。'),
    ('如何查看调用记录？', '登录 api.wf1.one →「日志」可查看每次调用的消耗、模型、时间等信息。'),
    ('支持流式输出吗？', '支持 ✅ 所有兼容 OpenAI 的流式调用均可正常使用（stream=true）。'),
    ('支持图片识别（Vision）吗？', '支持 ✅ 多模态模型（如 gpt-4o、claude-3.5-sonnet、gemini-pro-vision）通过标准的接口格式传入图片 URL 或 Base64 即可使用。'),
]

for q, a in faqs:
    p = doc.add_paragraph()
    p.add_run(f'❓ {q}').bold = True
    doc.add_paragraph(a)
    doc.add_paragraph()  # spacer

# Footer
add_separator(doc)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('生命不息，折腾不止 🚀')
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x0d, 0x47, 0xa1)
r.bold = True

# Save
output_path = '/opt/data/workspace/WF_AI中转站_大模型对接完整教程.docx'
doc.save(output_path)
print(f'✅ Word 文档已生成: {output_path}')
